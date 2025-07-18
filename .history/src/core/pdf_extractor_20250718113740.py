"""
PDF-based extraction module for SpecConverter v1.0.

This module provides an alternative extraction method using PDF conversion
and OCR to extract text with numbering from Word documents.
"""

import os
import tempfile
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import json

try:
    import pypandoc
    PYPANDOC_AVAILABLE = True
except ImportError:
    PYPANDOC_AVAILABLE = False
    logging.warning("pypandoc not available. PDF conversion will use alternative methods.")

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    logging.warning("PyPDF2 not available. PDF text extraction will be limited.")

try:
    import pytesseract
    from pdf2image import convert_from_path
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("OCR dependencies not available. OCR extraction will be disabled.")

from .models import ContentBlock, SpecDocument, HeaderFooterData, TemplateAnalysis, ValidationResults


class PDFExtractor:
    """
    Extract content from Word documents using PDF conversion and OCR.
    
    This approach converts Word documents to PDF, then uses OCR to extract
    text with numbering, which can be more reliable for complex formatting.
    """
    
    def __init__(self, temp_dir: Optional[str] = None):
        """
        Initialize the PDF extractor.
        
        Args:
            temp_dir: Directory for temporary files. If None, uses system temp.
        """
        self.temp_dir = temp_dir or tempfile.gettempdir()
        self.logger = logging.getLogger(__name__)
        
        # Check dependencies
        if not PYPANDOC_AVAILABLE:
            self.logger.warning("pypandoc not available. PDF conversion may fail.")
        if not OCR_AVAILABLE:
            self.logger.warning("OCR dependencies not available. OCR extraction disabled.")
    
    def extract_document(self, docx_path: str, template_path: Optional[str] = None) -> SpecDocument:
        """
        Extract content from a Word document using PDF conversion and OCR.
        
        Args:
            docx_path: Path to the Word document
            template_path: Optional template path (not used in PDF extraction)
            
        Returns:
            SpecDocument with extracted content
        """
        self.logger.info(f"Starting PDF-based extraction of {docx_path}")
        
        # Convert Word to PDF
        pdf_path = self._convert_docx_to_pdf(docx_path)
        if not pdf_path:
            raise ValueError(f"Failed to convert {docx_path} to PDF")
        
        try:
            # Extract text using multiple methods
            text_content = self._extract_text_from_pdf(pdf_path)
            
            # Parse content blocks from extracted text
            content_blocks = self._parse_content_blocks(text_content)
            
            # Create document structure
            document = SpecDocument(
                file_path=docx_path,
                content_blocks=content_blocks,
                header_footer=HeaderFooterData(
                    header={},
                    footer={},
                    margins={},
                    document_settings={}
                ),
                template_analysis=TemplateAnalysis(
                    template_path=template_path or "",
                    analysis_timestamp="",
                    numbering_definitions={},
                    bwa_list_levels={},
                    level_mappings={},
                    summary={}
                ),
                validation_results=ValidationResults(
                    errors=[],
                    corrections=[],
                    validation_summary={}
                )
            )
            
            self.logger.info(f"Successfully extracted {len(content_blocks)} content blocks")
            return document
            
        finally:
            # Clean up temporary PDF
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
    
    def _convert_docx_to_pdf(self, docx_path: str) -> Optional[str]:
        """
        Convert Word document to PDF using multiple methods.
        
        Args:
            docx_path: Path to Word document
            
        Returns:
            Path to generated PDF, or None if conversion failed
        """
        try:
            # Create temporary PDF path
            pdf_path = os.path.join(
                self.temp_dir, 
                f"temp_{Path(docx_path).stem}.pdf"
            )
            
            # Method 1: Try using python-docx to extract text directly
            self.logger.info("Attempting direct text extraction from Word document...")
            try:
                from docx import Document
                doc = Document(docx_path)
                text_content = []
                for paragraph in doc.paragraphs:
                    text_content.append(paragraph.text)
                
                # Save as text file for now (we'll process it as if it came from PDF)
                text_path = os.path.join(self.temp_dir, f"temp_{Path(docx_path).stem}.txt")
                with open(text_path, 'w', encoding='utf-8') as f:
                    f.write('\n'.join(text_content))
                
                self.logger.info(f"Text extracted successfully: {text_path}")
                return text_path  # Return text path instead of PDF for now
                
            except Exception as e:
                self.logger.warning(f"Direct text extraction failed: {e}")
            
            # Method 2: Try pandoc if available
            if PYPANDOC_AVAILABLE:
                try:
                    self.logger.info("Converting Word to PDF using pandoc...")
                    pypandoc.convert_file(
                        docx_path,
                        'pdf',
                        outputfile=pdf_path,
                        extra_args=['--pdf-engine=wkhtmltopdf']
                    )
                    if os.path.exists(pdf_path):
                        self.logger.info(f"PDF created successfully: {pdf_path}")
                        return pdf_path
                except Exception as e:
                    self.logger.warning(f"Pandoc conversion failed: {e}")
            
            # Method 3: Try LibreOffice if available
            try:
                self.logger.info("Attempting conversion using LibreOffice...")
                import subprocess
                result = subprocess.run([
                    'soffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', self.temp_dir, docx_path
                ], capture_output=True, text=True)
                
                if result.returncode == 0:
                    # LibreOffice creates PDF in the same directory as input
                    pdf_path = str(Path(docx_path).with_suffix('.pdf'))
                    if os.path.exists(pdf_path):
                        self.logger.info(f"PDF created successfully: {pdf_path}")
                        return pdf_path
            except Exception as e:
                self.logger.warning(f"LibreOffice conversion failed: {e}")
            
            self.logger.error("All PDF conversion methods failed")
            return None
                
        except Exception as e:
            self.logger.error(f"PDF conversion failed: {e}")
            return None
    
    def _extract_text_from_pdf(self, pdf_path: str) -> str:
        """
        Extract text from PDF or text file using multiple methods.
        
        Args:
            pdf_path: Path to PDF or text file
            
        Returns:
            Extracted text content
        """
        # Check if it's a text file
        if pdf_path.endswith('.txt'):
            self.logger.info("Reading text file directly...")
            with open(pdf_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        text_content = ""
        
        # Method 1: Try direct text extraction
        if PYPDF2_AVAILABLE:
            try:
                text_content = self._extract_text_pypdf2(pdf_path)
                if text_content.strip():
                    self.logger.info("Successfully extracted text using PyPDF2")
                    return text_content
            except Exception as e:
                self.logger.warning(f"PyPDF2 extraction failed: {e}")
        
        # Method 2: Use OCR if available
        if OCR_AVAILABLE:
            try:
                text_content = self._extract_text_ocr(pdf_path)
                if text_content.strip():
                    self.logger.info("Successfully extracted text using OCR")
                    return text_content
            except Exception as e:
                self.logger.warning(f"OCR extraction failed: {e}")
        
        # Method 3: Try pandoc PDF to text
        if PYPANDOC_AVAILABLE:
            try:
                text_content = self._extract_text_pandoc(pdf_path)
                if text_content.strip():
                    self.logger.info("Successfully extracted text using pandoc")
                    return text_content
            except Exception as e:
                self.logger.warning(f"Pandoc extraction failed: {e}")
        
        if not text_content.strip():
            raise ValueError("Failed to extract any text from PDF using all available methods")
        
        return text_content
    
    def _extract_text_pypdf2(self, pdf_path: str) -> str:
        """Extract text using PyPDF2."""
        text = ""
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def _extract_text_ocr(self, pdf_path: str) -> str:
        """Extract text using OCR (Tesseract)."""
        # Convert PDF to images
        images = convert_from_path(pdf_path, dpi=300)
        
        text = ""
        for i, image in enumerate(images):
            self.logger.info(f"Processing page {i+1}/{len(images)} with OCR...")
            
            # Extract text from image
            page_text = pytesseract.image_to_string(image, config='--psm 6')
            text += page_text + "\n"
        
        return text
    
    def _extract_text_pandoc(self, pdf_path: str) -> str:
        """Extract text using pandoc."""
        return pypandoc.convert_file(pdf_path, 'plain')
    
    def _parse_content_blocks(self, text_content: str) -> List[ContentBlock]:
        """
        Parse extracted text into content blocks.
        
        Args:
            text_content: Raw extracted text
            
        Returns:
            List of parsed content blocks
        """
        blocks = []
        lines = text_content.split('\n')
        
        current_block = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Try to identify section headers (numbered sections)
            if self._is_section_header(line):
                # Save previous block if exists
                if current_block and current_content:
                    current_block.content = '\n'.join(current_content)
                    blocks.append(current_block)
                
                # Start new block
                section_number = self._extract_section_number(line)
                current_block = ContentBlock(
                    text=line,
                    level_type="section",
                    number=section_number,
                    content="",
                    level_number=self._determine_level(section_number)
                )
                current_content = []
            else:
                # Add to current block content
                current_content.append(line)
        
        # Add final block
        if current_block and current_content:
            current_block.content = '\n'.join(current_content)
            blocks.append(current_block)
        
        return blocks
    
    def _is_section_header(self, line: str) -> bool:
        """Check if a line appears to be a section header."""
        # Look for patterns like "SECTION 26 05 00", "26.05.00", etc.
        import re
        
        patterns = [
            r'^SECTION\s+\d+[\-\s]\d+[\-\s]\d+',  # SECTION 26-05-00 or SECTION 26 05 00
            r'^\d+\.\d+\.\d+',              # 26.05.00
            r'^\d+[\-\s]\d+[\-\s]\d+',      # 26-05-00 or 26 05 00
            r'^\d+\.\d+',                   # 1.0, 2.0, etc.
            r'^\d+\.\d+\s+[A-Z]',           # 1.01 DESCRIPTION, 2.01 PRODUCTS, etc.
            r'^[A-Z]\.\s+',                 # A., B., C., etc.
            r'^\d+\.\s+',                   # 1., 2., 3., etc.
            r'^[a-z]\.\s+',                 # a., b., c., etc.
            r'^PART\s+\d+',                 # PART 1
            r'^DIVISION\s+\d+',             # DIVISION 26
        ]
        
        for pattern in patterns:
            if re.match(pattern, line, re.IGNORECASE):
                return True
        
        return False
    
    def _extract_section_number(self, line: str) -> str:
        """Extract section number from header line."""
        import re
        
        # Try different patterns
        patterns = [
            r'SECTION\s+(\d+[\-\s]\d+[\-\s]\d+)',
            r'(\d+\.\d+\.\d+)',
            r'(\d+[\-\s]\d+[\-\s]\d+)',
            r'(\d+\.\d+)',                   # 1.0, 2.0, etc.
            r'(\d+\.\d+)\s+[A-Z]',           # 1.01 DESCRIPTION -> 1.01
            r'([A-Z])\.\s+',                 # A. -> A
            r'(\d+)\.\s+',                   # 1. -> 1
            r'([a-z])\.\s+',                 # a. -> a
            r'PART\s+(\d+)',
            r'DIVISION\s+(\d+)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, line, re.IGNORECASE)
            if match:
                return match.group(1)
        
        # Fallback: return the whole line
        return line
    
    def _determine_level(self, section_number: str) -> int:
        """Determine the hierarchical level of a section."""
        if 'DIVISION' in section_number.upper():
            return 0
        elif 'PART' in section_number.upper():
            return 1
        elif section_number.count('.') == 2 or section_number.count('-') == 2:
            return 0  # Main section like 26.05.00 or 26-05-00
        elif section_number.count('.') == 1:
            if section_number.endswith('.0'):
                return 0  # 1.0, 2.0, etc.
            else:
                return 1  # 1.01, 2.01, etc.
        elif section_number.isupper() and len(section_number) == 1:
            return 2  # A, B, C, etc.
        elif section_number.isdigit():
            return 3  # 1, 2, 3, etc.
        elif section_number.islower() and len(section_number) == 1:
            return 4  # a, b, c, etc.
        else:
            return 5  # Default level


def extract_via_pdf(docx_path: str, template_path: Optional[str] = None) -> SpecDocument:
    """
    Convenience function to extract content using PDF conversion and OCR.
    
    Args:
        docx_path: Path to Word document
        template_path: Optional template path
        
    Returns:
        Extracted SpecDocument
    """
    extractor = PDFExtractor()
    return extractor.extract_document(docx_path, template_path) 