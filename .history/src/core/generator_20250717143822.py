from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.oxml.ns import qn
import json
import os
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.enum.style import WD_STYLE_TYPE
from typing import Dict, List, Any, Optional

# Configuration variables - change these to modify font and size for all text
TEMPLATE_PATH = '../templates/test_template_RPA.docx'
OUTPUT_PATH   = '../output/210500 Common Work Results For Fire Suppression.docx'
CONTENT_PATH  = '../output/210500 Common Work Results For Fire Suppression_v3.json'
FONT_NAME = 'Arial'
FONT_SIZE = 10

def list_available_styles(doc: Any) -> None:
    """List all available styles in the document"""
    print("DEBUG: Available styles in template:")
    for style in doc.styles:
        print(f"  - {style.name}")

def clear_document(doc: Any) -> None:
    # Remove all paragraphs
    for _ in range(len(doc.paragraphs)):
        p = doc.paragraphs[0]
        p._element.getparent().remove(p._element)
    # Remove all tables
    for _ in range(len(doc.tables)):
        t = doc.tables[0]
        t._element.getparent().remove(t._element)

def set_font_and_size(paragraph: Any) -> None:
    """Set font and size for all runs in a paragraph"""
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE)
        # For compatibility with some versions of Word
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

def apply_styling_from_json(paragraph: Any, block: Dict[str, Any]) -> None:
    """Apply styling information from JSON block to paragraph"""
    try:
        # Font properties
        if block.get('font_name'):
            for run in paragraph.runs:
                run.font.name = block['font_name']
                # For compatibility with some versions of Word
                r = run._element
                if r.rPr is None:
                    r_pr = OxmlElement('w:rPr')
                    r.insert(0, r_pr)
                r_fonts = r.rPr.find(qn('w:rFonts'))
                if r_fonts is None:
                    r_fonts = OxmlElement('w:rFonts')
                    r.rPr.append(r_fonts)
                r_fonts.set(qn('w:eastAsia'), block['font_name'])
        
        if block.get('font_size'):
            for run in paragraph.runs:
                run.font.size = Pt(block['font_size'])
        
        if block.get('font_bold') is not None:
            for run in paragraph.runs:
                run.font.bold = block['font_bold']
        
        if block.get('font_italic') is not None:
            for run in paragraph.runs:
                run.font.italic = block['font_italic']
        
        if block.get('font_underline'):
            for run in paragraph.runs:
                # Map underline values to valid WD_UNDERLINE constants
                underline_map = {
                    'single': WD_UNDERLINE.SINGLE,
                    'double': WD_UNDERLINE.DOUBLE,
                    'thick': WD_UNDERLINE.THICK,
                    'dotted': WD_UNDERLINE.DOTTED,
                    'dash': WD_UNDERLINE.DASH,
                    'dotDash': WD_UNDERLINE.DOT_DASH,
                    'dotDotDash': WD_UNDERLINE.DOT_DOT_DASH,
                    'wavy': WD_UNDERLINE.WAVY,
                    'none': WD_UNDERLINE.NONE
                }
                underline_value = underline_map.get(block['font_underline'], WD_UNDERLINE.SINGLE)
                run.font.underline = underline_value
                r = run._element
                if r.rPr is None:
                    r_pr = OxmlElement('w:rPr')
                    r.insert(0, r_pr)
        
        if block.get('font_color'):
            for run in paragraph.runs:
                color_val = block['font_color']
                try:
                    run.font.color.rgb = RGBColor.from_string(color_val)
                except Exception:
                    # Default to black if invalid
                    run.font.color.rgb = RGBColor.from_string('000000')
        
        # Additional font properties
        if block.get('font_strike') is not None:
            for run in paragraph.runs:
                run.font.strike = block['font_strike']
        
        if block.get('font_small_caps') is not None:
            for run in paragraph.runs:
                run.font.small_caps = block['font_small_caps']
        
        if block.get('font_all_caps') is not None:
            for run in paragraph.runs:
                run.font.all_caps = block['font_all_caps']
        
        # Paragraph properties
        if block.get('paragraph_alignment'):
            alignment_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'both': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            if block['paragraph_alignment'] in alignment_map:
                paragraph.alignment = alignment_map[block['paragraph_alignment']]
        
        # Indentation (convert points to inches for Word)
        if block.get('paragraph_indent_left'):
            paragraph.paragraph_format.left_indent = Inches(block['paragraph_indent_left'] / 72.0)
        
        if block.get('paragraph_indent_right'):
            paragraph.paragraph_format.right_indent = Inches(block['paragraph_indent_right'] / 72.0)
        
        if block.get('paragraph_indent_first_line'):
            paragraph.paragraph_format.first_line_indent = Inches(block['paragraph_indent_first_line'] / 72.0)
        
        # Spacing
        if block.get('paragraph_spacing_before'):
            paragraph.paragraph_format.space_before = Pt(block['paragraph_spacing_before'])
        
        if block.get('paragraph_spacing_after'):
            paragraph.paragraph_format.space_after = Pt(block['paragraph_spacing_after'])
        
        if block.get('paragraph_line_spacing'):
            paragraph.paragraph_format.line_spacing = block['paragraph_line_spacing']
        
        # Additional paragraph properties
        if block.get('paragraph_keep_with_next') is not None:
            paragraph.paragraph_format.keep_with_next = block['paragraph_keep_with_next']
        
        if block.get('paragraph_keep_lines_together') is not None:
            paragraph.paragraph_format.keep_lines_together = block['paragraph_keep_lines_together']
        
        if block.get('paragraph_page_break_before') is not None:
            paragraph.paragraph_format.page_break_before = block['paragraph_page_break_before']
        
        if block.get('paragraph_widow_control') is not None:
            paragraph.paragraph_format.widow_control = block['paragraph_widow_control']
        
        # Don't add space between paragraphs of the same style
        if block.get('paragraph_dont_add_space_between_same_style') is not None:
            paragraph.paragraph_format.dont_add_space_between_same_style = block['paragraph_dont_add_space_between_same_style']
        
    except Exception as e:
        print(f"Warning: Could not apply styling from JSON: {e}")
        # Fallback to default styling
        set_font_and_size(paragraph)

    # Highlight corrected blocks (used_fallback_styling) in yellow
    if block.get('used_fallback_styling'):
        for run in paragraph.runs:
            r = run._element
            if r.rPr is None:
                r_pr = OxmlElement('w:rPr')
                r.insert(0, r_pr)
            highlight = OxmlElement('w:highlight')
            highlight.set(qn('w:val'), 'yellow')
            r.rPr.append(highlight)

def apply_style_definitions_from_json(doc: Any, json_data: Optional[Dict[str, Any]]) -> None:
    """Apply style definitions from JSON to ensure proper styling in regenerated document"""
    try:
        if json_data is None:
            return
        content_blocks = json_data.get('content_blocks', [])
        if not content_blocks:
            return
        
        # Collect all unique BWA level names used in content blocks
        bwa_style_names = set()
        for block in content_blocks:
            bwa_level_name = block.get('bwa_level_name')
            if bwa_level_name and bwa_level_name != 'Normal':
                bwa_style_names.add(bwa_level_name)
        
        # Apply styling for each BWA style
        for bwa_style_name in bwa_style_names:
            apply_style_definition(doc, bwa_style_name, content_blocks)
        
        print(f"Applied style definitions for {len(bwa_style_names)} BWA styles")
        
    except Exception as e:
        print(f"Warning: Could not apply style definitions: {e}")

def apply_style_definition(doc: Any, style_name: str, content_blocks: List[Dict[str, Any]]) -> None:
    """Apply definition for a specific style based on content blocks using that style"""
    try:
        # Get or create the style
        if style_name in doc.styles:
            style = doc.styles[style_name]
        else:
            # Create new style if it doesn't exist
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        
        # Find content blocks using this BWA level name to determine its properties
        style_blocks = [block for block in content_blocks if block.get('bwa_level_name') == style_name]
        if not style_blocks:
            return
        
        # Use the first block to determine style properties (they should be consistent)
        sample_block = style_blocks[0]
        
        # Apply paragraph format properties
        if hasattr(style, 'paragraph_format') and style.paragraph_format:
            pf = style.paragraph_format
            
            # Alignment
            if sample_block.get('paragraph_alignment'):
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                alignment_map = {
                    'left': WD_ALIGN_PARAGRAPH.LEFT,
                    'center': WD_ALIGN_PARAGRAPH.CENTER,
                    'right': WD_ALIGN_PARAGRAPH.RIGHT,
                    'both': WD_ALIGN_PARAGRAPH.JUSTIFY
                }
                if sample_block['paragraph_alignment'] in alignment_map:
                    pf.alignment = alignment_map[sample_block['paragraph_alignment']]
            
            # Indentation
            if sample_block.get('paragraph_indent_left'):
                from docx.shared import Inches
                pf.left_indent = Inches(sample_block['paragraph_indent_left'] / 72.0)
            
            if sample_block.get('paragraph_indent_right'):
                from docx.shared import Inches
                pf.right_indent = Inches(sample_block['paragraph_indent_right'] / 72.0)
            
            if sample_block.get('paragraph_indent_first_line'):
                from docx.shared import Inches
                pf.first_line_indent = Inches(sample_block['paragraph_indent_first_line'] / 72.0)
            
            # Spacing
            if sample_block.get('paragraph_spacing_before'):
                from docx.shared import Pt
                pf.space_before = Pt(sample_block['paragraph_spacing_before'])
            
            if sample_block.get('paragraph_spacing_after'):
                from docx.shared import Pt
                pf.space_after = Pt(sample_block['paragraph_spacing_after'])
            
            if sample_block.get('paragraph_line_spacing'):
                pf.line_spacing = sample_block['paragraph_line_spacing']
            
            # Other paragraph properties
            if sample_block.get('paragraph_keep_with_next') is not None:
                pf.keep_with_next = sample_block['paragraph_keep_with_next']
            
            if sample_block.get('paragraph_keep_lines_together') is not None:
                pf.keep_lines_together = sample_block['paragraph_keep_lines_together']
            
            if sample_block.get('paragraph_page_break_before') is not None:
                pf.page_break_before = sample_block['paragraph_page_break_before']
            
            if sample_block.get('paragraph_widow_control') is not None:
                pf.widow_control = sample_block['paragraph_widow_control']
        
        # Apply font properties
        if hasattr(style, 'font') and style.font:
            font = style.font
            
            if sample_block.get('font_name'):
                font.name = sample_block['font_name']
            
            if sample_block.get('font_size'):
                from docx.shared import Pt
                font.size = Pt(sample_block['font_size'])
            
            if sample_block.get('font_bold') is not None:
                font.bold = sample_block['font_bold']
            
            if sample_block.get('font_italic') is not None:
                font.italic = sample_block['font_italic']
            
            if sample_block.get('font_underline'):
                from docx.enum.text import WD_UNDERLINE
                underline_map = {
                    'single': WD_UNDERLINE.SINGLE,
                    'double': WD_UNDERLINE.DOUBLE,
                    'thick': WD_UNDERLINE.THICK,
                    'dotted': WD_UNDERLINE.DOTTED,
                    'dash': WD_UNDERLINE.DASH,
                    'dotDash': WD_UNDERLINE.DOT_DASH,
                    'dotDotDash': WD_UNDERLINE.DOT_DOT_DASH,
                    'wavy': WD_UNDERLINE.WAVY,
                    'none': WD_UNDERLINE.NONE
                }
                if sample_block['font_underline'] in underline_map:
                    font.underline = underline_map[sample_block['font_underline']]
            
            if sample_block.get('font_color'):
                from docx.shared import RGBColor
                font.color.rgb = RGBColor.from_string(sample_block['font_color'])
            
            if sample_block.get('font_strike') is not None:
                font.strike = sample_block['font_strike']
            
            if sample_block.get('font_small_caps') is not None:
                font.small_caps = sample_block['font_small_caps']
            
            if sample_block.get('font_all_caps') is not None:
                font.all_caps = sample_block['font_all_caps']
        
    except Exception as e:
        print(f"Warning: Could not apply style definition for {style_name}: {e}")

def apply_document_settings_from_json(doc: Any, json_data: Optional[Dict[str, Any]]) -> None:
    """Apply document-level settings from JSON to the document"""
    try:
        if json_data is None:
            return
        document_settings = json_data.get('document_settings', {})
        if not document_settings:
            return
        
        # Apply settings to the first section
        section = doc.sections[0]
        
        # Page size and orientation
        if document_settings.get('page_width') and document_settings.get('page_height'):
            from docx.shared import Inches
            section.page_width = Inches(document_settings['page_width'])
            section.page_height = Inches(document_settings['page_height'])
        
        # Margins (if not already set by template)
        if document_settings.get('top_margin'):
            section.top_margin = Inches(document_settings['top_margin'])
        if document_settings.get('bottom_margin'):
            section.bottom_margin = Inches(document_settings['bottom_margin'])
        if document_settings.get('left_margin'):
            section.left_margin = Inches(document_settings['left_margin'])
        if document_settings.get('right_margin'):
            section.right_margin = Inches(document_settings['right_margin'])
        
        # Header and footer distances
        if document_settings.get('header_distance'):
            section.header_distance = Inches(document_settings['header_distance'])
        if document_settings.get('footer_distance'):
            section.footer_distance = Inches(document_settings['footer_distance'])
        
        # Gutter settings
        if document_settings.get('gutter'):
            section.gutter = Inches(document_settings['gutter'])
        
        # Different first page header/footer
        if document_settings.get('different_first_page_header_footer') is not None:
            section.different_first_page_header_footer = document_settings['different_first_page_header_footer']
        
        # Different odd and even pages
        if document_settings.get('different_odd_and_even_pages') is not None:
            section.different_odd_and_even_pages = document_settings['different_odd_and_even_pages']
        
        # Page numbering
        if document_settings.get('page_numbering'):
            page_num = document_settings['page_numbering']
            if hasattr(section, 'page_numbering') and section.page_numbering:
                if page_num.get('start') is not None:
                    section.page_numbering.start = page_num['start']
                if page_num.get('restart') is not None:
                    section.page_numbering.restart = page_num['restart']
                if page_num.get('format') is not None:
                    section.page_numbering.format = page_num['format']
        
        # Line numbering
        if document_settings.get('line_numbering'):
            line_num = document_settings['line_numbering']
            if hasattr(section, 'line_numbering') and section.line_numbering:
                if line_num.get('start') is not None:
                    section.line_numbering.start = line_num['start']
                if line_num.get('increment') is not None:
                    section.line_numbering.increment = line_num['increment']
                if line_num.get('restart') is not None:
                    section.line_numbering.restart = line_num['restart']
                if line_num.get('distance') is not None:
                    section.line_numbering.distance = Inches(line_num['distance'])
        
        # Document properties
        if document_settings.get('document_properties') and doc.core_properties:
            props = document_settings['document_properties']
            if props.get('title'):
                doc.core_properties.title = props['title']
            if props.get('subject'):
                doc.core_properties.subject = props['subject']
            if props.get('author'):
                doc.core_properties.author = props['author']
            if props.get('keywords'):
                doc.core_properties.keywords = props['keywords']
            if props.get('category'):
                doc.core_properties.category = props['category']
            if props.get('comments'):
                doc.core_properties.comments = props['comments']
            if props.get('last_modified_by'):
                doc.core_properties.last_modified_by = props['last_modified_by']
            if props.get('revision'):
                doc.core_properties.revision = props['revision']
        
        # Apply default formatting settings
        if document_settings.get('default_formatting'):
            apply_default_formatting_from_json(doc, document_settings['default_formatting'])
        
        # Apply document-wide settings
        if document_settings.get('document_wide_settings'):
            apply_document_wide_settings_from_json(doc, document_settings['document_wide_settings'])
        
    except Exception as e:
        print(f"Warning: Could not apply document settings from JSON: {e}")

def apply_default_formatting_from_json(doc: Any, default_formatting: Dict[str, Any]) -> None:
    """Apply default formatting settings from JSON to document styles"""
    try:
        # Get the Normal style
        normal_style = doc.styles['Normal'] if 'Normal' in doc.styles else None
        if not normal_style:
            return
        
        # Apply default paragraph format
        if default_formatting.get('default_paragraph_format'):
            pf_data = default_formatting['default_paragraph_format']
            pf = normal_style.paragraph_format
            
            if pf_data.get('alignment'):
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                alignment_map = {
                    'left': WD_ALIGN_PARAGRAPH.LEFT,
                    'center': WD_ALIGN_PARAGRAPH.CENTER,
                    'right': WD_ALIGN_PARAGRAPH.RIGHT,
                    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                }
                if pf_data['alignment'] in alignment_map:
                    pf.alignment = alignment_map[pf_data['alignment']]
            
            if pf_data.get('left_indent'):
                from docx.shared import Inches
                pf.left_indent = Inches(pf_data['left_indent'])
            
            if pf_data.get('right_indent'):
                from docx.shared import Inches
                pf.right_indent = Inches(pf_data['right_indent'])
            
            if pf_data.get('first_line_indent'):
                from docx.shared import Inches
                pf.first_line_indent = Inches(pf_data['first_line_indent'])
            
            if pf_data.get('space_before'):
                from docx.shared import Pt
                pf.space_before = Pt(pf_data['space_before'])
            
            if pf_data.get('space_after'):
                from docx.shared import Pt
                pf.space_after = Pt(pf_data['space_after'])
            
            if pf_data.get('line_spacing'):
                pf.line_spacing = pf_data['line_spacing']
            
            if pf_data.get('keep_with_next') is not None:
                pf.keep_with_next = pf_data['keep_with_next']
            
            if pf_data.get('keep_lines_together') is not None:
                pf.keep_lines_together = pf_data['keep_lines_together']
            
            if pf_data.get('page_break_before') is not None:
                pf.page_break_before = pf_data['page_break_before']
            
            if pf_data.get('widow_control') is not None:
                pf.widow_control = pf_data['widow_control']
        
        # Apply default run format
        if default_formatting.get('default_run_format'):
            font_data = default_formatting['default_run_format']
            font = normal_style.font
            
            if font_data.get('name'):
                font.name = font_data['name']
            
            if font_data.get('size'):
                from docx.shared import Pt
                font.size = Pt(font_data['size'])
            
            if font_data.get('bold') is not None:
                font.bold = font_data['bold']
            
            if font_data.get('italic') is not None:
                font.italic = font_data['italic']
            
            if font_data.get('underline'):
                from docx.enum.text import WD_UNDERLINE
                underline_map = {
                    'single': WD_UNDERLINE.SINGLE,
                    'double': WD_UNDERLINE.DOUBLE,
                    'thick': WD_UNDERLINE.THICK,
                    'dotted': WD_UNDERLINE.DOTTED,
                    'dash': WD_UNDERLINE.DASH,
                    'dotDash': WD_UNDERLINE.DOT_DASH,
                    'dotDotDash': WD_UNDERLINE.DOT_DOT_DASH,
                    'wavy': WD_UNDERLINE.WAVY,
                    'none': WD_UNDERLINE.NONE
                }
                if font_data['underline'] in underline_map:
                    font.underline = underline_map[font_data['underline']]
            
            if font_data.get('color'):
                from docx.shared import RGBColor
                font.color.rgb = RGBColor.from_string(font_data['color'])
            
            if font_data.get('strike') is not None:
                font.strike = font_data['strike']
            
            if font_data.get('small_caps') is not None:
                font.small_caps = font_data['small_caps']
            
            if font_data.get('all_caps') is not None:
                font.all_caps = font_data['all_caps']
        
    except Exception as e:
        print(f"Warning: Could not apply default formatting from JSON: {e}")

def apply_document_wide_settings_from_json(doc: Any, doc_wide_settings: Dict[str, Any]) -> None:
    """Apply document-wide settings from JSON to document"""
    try:
        # Note: Most document-wide settings are read-only or require direct XML manipulation
        # We'll focus on settings that can be applied through the python-docx API
        
        # Default tab stop can be applied through styles
        if doc_wide_settings.get('default_tab_stop'):
            # Apply to Normal style paragraph format
            normal_style = doc.styles.get('Normal')
            if normal_style and normal_style.paragraph_format:
                # Convert twips to inches (1 inch = 1440 twips)
                tab_stop_inches = float(doc_wide_settings['default_tab_stop']) / 1440.0
                from docx.shared import Inches
                normal_style.paragraph_format.tab_stops.add_tab_stop(Inches(tab_stop_inches))
        
        # Track changes setting
        if doc_wide_settings.get('track_changes') is not None:
            # This would require direct XML manipulation
            print(f"DEBUG: Track changes setting found: {doc_wide_settings['track_changes']} (requires XML manipulation)")
        
        # Document protection
        if doc_wide_settings.get('document_protection'):
            protection = doc_wide_settings['document_protection']
            print(f"DEBUG: Document protection settings found: {protection} (requires XML manipulation)")
        
    except Exception as e:
        print(f"Warning: Could not apply document-wide settings from JSON: {e}")

def apply_explicit_indentation(paragraph: Any, bwa_level_name: Optional[str], level_number: Optional[int], json_data: Optional[Dict[str, Any]]) -> None:
    """
    Apply explicit indentation based on template analysis to override Word's default indentation
    
    Args:
        paragraph: The paragraph to apply indentation to
        bwa_level_name: The BWA level name (e.g., "BWA-Item", "BWA-List")
        level_number: The level number (0, 1, 2, 3, etc.)
        json_data: The JSON data containing template analysis
    """
    try:
        if not json_data or not bwa_level_name or level_number is None:
            return
        
        template_analysis = json_data.get('template_analysis', {})
        if not template_analysis:
            return
        
        # Get the template numbering definitions
        numbering_definitions = template_analysis.get('template_numbering', {})
        if not numbering_definitions:
            return
        
        # Find the correct abstract numbering definition (abstract_num_id "1" for BWA styles)
        # This maps to numId 10-25 in the template
        abstract_num_id = "1"  # This is the one with BWA styles
        
        if abstract_num_id in numbering_definitions:
            abstract_info = numbering_definitions[abstract_num_id]
            levels = abstract_info.get('levels', {})
            
            # Get the level info for this specific level
            level_str = str(level_number)
            if level_str in levels:
                level_info = levels[level_str]
                p_pr = level_info.get('pPr', {})
                
                # Extract indentation values (convert from twips to inches)
                if 'indent' in p_pr:
                    indent = p_pr['indent']
                    
                    # Left indent
                    if indent.get('left'):
                        left_twips = float(indent['left'])
                        left_inches = left_twips / 1440.0  # Convert twips to inches
                        paragraph.paragraph_format.left_indent = Inches(left_inches)
                        print(f"DEBUG: Applied left indent {left_inches:.3f} inches for {bwa_level_name} level {level_number}")
                    
                    # First line indent (hanging indent)
                    if indent.get('firstLine'):
                        first_line_twips = float(indent['firstLine'])
                        first_line_inches = first_line_twips / 1440.0
                        paragraph.paragraph_format.first_line_indent = Inches(first_line_inches)
                        print(f"DEBUG: Applied first line indent {first_line_inches:.3f} inches for {bwa_level_name} level {level_number}")
                    
                    # Hanging indent
                    if indent.get('hanging'):
                        hanging_twips = float(indent['hanging'])
                        hanging_inches = hanging_twips / 1440.0
                        paragraph.paragraph_format.first_line_indent = Inches(-hanging_inches)  # Negative for hanging
                        print(f"DEBUG: Applied hanging indent {hanging_inches:.3f} inches for {bwa_level_name} level {level_number}")
                    
                    # Right indent
                    if indent.get('right'):
                        right_twips = float(indent['right'])
                        right_inches = right_twips / 1440.0
                        paragraph.paragraph_format.right_indent = Inches(right_inches)
                        print(f"DEBUG: Applied right indent {right_inches:.3f} inches for {bwa_level_name} level {level_number}")
        
    except Exception as e:
        print(f"Warning: Could not apply explicit indentation: {e}")

def apply_margins_from_json(doc: Any, json_data: Optional[Dict[str, Any]]) -> None:
    """Apply margin settings from JSON to the document"""
    try:
        if json_data is None:
            return
        margins = json_data.get('margins', {})
        if not margins:
            return
        
        # Apply margins to the first section
        section = doc.sections[0]
        
        if margins.get('top_margin'):
            section.top_margin = Inches(margins['top_margin'])
        if margins.get('bottom_margin'):
            section.bottom_margin = Inches(margins['bottom_margin'])
        if margins.get('left_margin'):
            section.left_margin = Inches(margins['left_margin'])
        if margins.get('right_margin'):
            section.right_margin = Inches(margins['right_margin'])
        if margins.get('header_distance'):
            section.header_distance = Inches(margins['header_distance'])
        if margins.get('footer_distance'):
            section.footer_distance = Inches(margins['footer_distance'])
        
    except Exception as e:
        print(f"Warning: Could not apply margins from JSON: {e}")

def parse_spec_json(json_path: str) -> Optional[Dict[str, Any]]:
    """Parse JSON file and return structured data"""
    
    if not os.path.exists(json_path):
        print(f"ERROR: JSON file not found: {json_path}")
        return None
    
    try:
        with open(json_path, 'r', encoding='utf-8') as file:
            content = json.load(file)
            return content
    except json.JSONDecodeError as e:
        print(f"ERROR: Invalid JSON format: {e}")
        return None
    except Exception as e:
        print(f"ERROR: Failed to read JSON file: {e}")
        return None

def get_style_for_bwa_level(bwa_level_name: str) -> str:
    """Map BWA level names to template style names"""
    style_mapping = {
        "BWA-SectionNumber": "BWA-SectionNumber",
        "BWA-SectionTitle": "BWA-SectionTitle",
        "BWA-PART": "BWA-PART",
        "BWA-SUBSECTION": "BWA-SUBSECTION", 
        "BWA-Item": "BWA-Item",
        "BWA-List": "BWA-List",
        "BWA-SubList": "BWA-SubList",
        "BWA-SubItem": "BWA-SubItem",
        "BWA-SubSubItem": "BWA-SubSubItem",
        "BWA-SubSubList": "BWA-SubSubList"
    }
    return style_mapping.get(bwa_level_name, "Normal")

def clean_text_for_display(text: str, level_type: str, number: Optional[str]) -> str:
    """Clean text by removing numbering prefixes while preserving content"""
    if not text:
        return text
    
    # Remove common numbering patterns from the beginning of text
    import re
    
    # Patterns to remove (in order of specificity)
    patterns = [
        # Item patterns: "A.\t", "B.\t", "C.\t", etc.
        r'^[A-Z]\.\s*\t\s*',
        # List patterns: "1.\t", "2.\t", "3.\t", etc.
        r'^\d+\.\s*\t\s*',
        # Sub-list patterns: "a.\t", "b.\t", "c.\t", etc.
        r'^[a-z]\.\s*\t\s*',
        # Part patterns: "1.0\t", "2.0\t", etc.
        r'^\d+\.0\s*\t\s*',
        # Subsection patterns: "1.01\t", "1.02\t", etc.
        r'^\d+\.\d{2}\s*\t\s*',
        # Alternative subsection patterns: "1.1\t", "1.2\t", etc.
        r'^\d+\.\d\s*\t\s*',
        # Section patterns: "SECTION 26 05 00\t"
        r'^SECTION\s+[^\t]*\s*\t\s*',
        # Generic tab removal at start
        r'^\s*\t\s*'
    ]
    
    cleaned_text = text
    for pattern in patterns:
        cleaned_text = re.sub(pattern, '', cleaned_text, flags=re.IGNORECASE)
    
    # If we removed everything, return the original text
    if not cleaned_text.strip():
        return text
    
    return cleaned_text.strip()

def update_numbering_context(numbering_context: Dict[str, int], level_type: str) -> str:
    """
    Update numbering context and return the correct sequential number for the level type.
    This ensures proper sequential numbering that restarts at appropriate levels.
    """
    if level_type == "part":
        numbering_context['part'] += 1
        # Reset lower-level counters when starting a new part
        numbering_context['subsection'] = 0
        numbering_context['item'] = 0
        numbering_context['list'] = 0
        numbering_context['sub_list'] = 0
        return f"{numbering_context['part']}.0"
    
    elif level_type == "subsection":
        numbering_context['subsection'] += 1
        # Reset lower-level counters when starting a new subsection
        numbering_context['item'] = 0
        numbering_context['list'] = 0
        numbering_context['sub_list'] = 0
        return f"{numbering_context['part']}.{numbering_context['subsection']:02d}"
    
    elif level_type == "item":
        numbering_context['item'] += 1
        # Reset lower-level counters when starting a new item
        numbering_context['list'] = 0
        numbering_context['sub_list'] = 0
        # Convert to letter (1=A, 2=B, 3=C, etc.)
        return chr(64 + numbering_context['item'])  # ASCII 65 = 'A'
    
    elif level_type == "list":
        numbering_context['list'] += 1
        # Reset sub-list counter when starting a new list
        numbering_context['sub_list'] = 0
        return str(numbering_context['list'])
    
    elif level_type == "sub_list":
        numbering_context['sub_list'] += 1
        # Convert to lowercase letter (1=a, 2=b, 3=c, etc.)
        return chr(96 + numbering_context['sub_list'])  # ASCII 97 = 'a'
    
    else:
        # For other level types, don't update counters
        return ""

def generate_content_from_v3_json(doc: Any, json_data: Optional[Dict[str, Any]]) -> None:
    """Generate document content from v3 JSON data using template styles"""
    
    if json_data is None:
        print("ERROR: No JSON data to process")
        return
    
    content_blocks = json_data.get('content_blocks', [])
    
    # Initialize numbering context for proper sequential numbering
    numbering_context = {
        'part': 0,
        'subsection': 0, 
        'item': 0,
        'list': 0,
        'sub_list': 0
    }
    
    for i, block in enumerate(content_blocks):
        text = block.get('text', '')
        content = block.get('content', '')  # This is the cleaned content without numbering prefixes
        level_type = block.get('level_type', 'content')
        bwa_level_name = block.get('bwa_level_name')
        original_number = block.get('number')
        level_number = block.get('level_number')
        
        # Skip empty content
        if not text.strip():
            continue
        
        # Generate correct sequential numbering based on level type and context
        correct_number = update_numbering_context(numbering_context, level_type)
        
        # Debug output for numbering changes
        if level_type in ["part", "subsection", "item"] and correct_number:
            print(f"DEBUG: {level_type.upper()} numbering: {original_number} -> {correct_number}")
        
        # Use the cleaned content from the JSON (already has numbering prefixes stripped)
        # Fallback to cleaning the text if content is empty
        display_text = content if content.strip() else clean_text_for_display(text, level_type, original_number)
        
        # Determine the style to use
        if bwa_level_name:
            style_name = get_style_for_bwa_level(bwa_level_name)
        else:
            # Map level types to BWA styles first, then fallback to Word styles
            if level_type == "section":
                style_name = "BWA-SectionNumber"  # Try BWA style first
            elif level_type == "title":
                style_name = "BWA-SectionTitle"   # Try BWA style first
            elif level_type == "part_title":
                style_name = "BWA-PART"           # Use BWA style
            elif level_type == "subsection_title":
                style_name = "BWA-SUBSECTION"     # Use BWA style
            else:
                style_name = "Normal"
        
        # Add paragraph with appropriate style using cleaned text
        try:
            paragraph = doc.add_paragraph(display_text, style=style_name)
            apply_styling_from_json(paragraph, block)
            
            # ENHANCED: Apply explicit indentation based on template analysis
            apply_explicit_indentation(paragraph, bwa_level_name, level_number, json_data)
            
            # Apply list numbering if level_number is specified
            if level_number is not None and level_type in ["list", "sub_list", "item"]:
                # Apply list numbering based on level_number
                # This will use the template's multilevel list definitions
                try:
                    # Set the list level (0-based in python-docx)
                    paragraph._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level_number
                    
                    # Use the correct numbering definition from the template
                    # numId 10-25 map to abstract_num_id "1" which has the correct BWA level configurations
                    # numId 1 maps to abstract_num_id "3" which has decimal format for level 2
                    # We want abstract_num_id "1" for proper BWA-Item (A, B, C) formatting
                    paragraph._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = 10
                    
                    # For BWA-Item levels, we need to preserve the extracted letter numbers (A, B, C, etc.)
                    # rather than letting Word generate sequential numbers (1, 2, 3, etc.)
                    if level_type == "item" and correct_number and correct_number.isalpha():
                        # The template's multilevel list should be configured to use upperLetter format
                        # for level 2 (BWA-Item), so Word should automatically generate A, B, C, etc.
                        # If it's not working, we may need to manually set the numbering format
                        print(f"DEBUG: Applied list numbering for item '{correct_number}' at level {level_number} using numId 10")
                    
                except Exception as e:
                    print(f"Warning: Could not apply list numbering to paragraph: {e}")
            
        except Exception as e:
            # Fallback to Normal style if the specified style is not found
            paragraph = doc.add_paragraph(display_text, style="Normal")
            apply_styling_from_json(paragraph, block)
            apply_explicit_indentation(paragraph, bwa_level_name, level_number, json_data)
        
        # Add some spacing for better readability
        if level_type in ["section", "title", "part_title"]:
            doc.add_paragraph()  # Add blank line after major sections

def check_template_styles(template_path: str) -> None:
    """Check template styles and provide feedback about style definitions"""
    try:
        doc = Document(template_path)
        print(f"\nTemplate Style Analysis for: {template_path}")
        print("-" * 50)
        
        # Check for BWA styles specifically
        bwa_styles = []
        other_styles = []
        
        for style in doc.styles:
            if style.name and style.name.startswith('BWA-'):
                bwa_styles.append(style.name)
            elif style.name and style.name not in ['Normal', 'Default Paragraph Font', 'Default Paragraph Font (Asian)', 'Default Paragraph Font (Complex Script)']:
                other_styles.append(style.name)
        
        print(f"BWA Styles found: {len(bwa_styles)}")
        for style_name in bwa_styles:
            print(f"  - {style_name}")
        
        if other_styles:
            print(f"\nOther custom styles found: {len(other_styles)}")
            for style_name in other_styles[:10]:  # Show first 10
                print(f"  - {style_name}")
            if len(other_styles) > 10:
                print(f"  ... and {len(other_styles) - 10} more")
        
        print(f"\nTotal styles in template: {len(doc.styles)}")
        print("\nNote: If styles appear with default formatting in regenerated documents,")
        print("consider setting them to 'New documents based on this template' in Word.")
        
    except Exception as e:
        print(f"Warning: Could not analyze template styles: {e}")

def clone_header_footer_styles(template_doc: Any, target_doc: Any) -> None:
    """Clone header and footer styles from template to target document"""
    try:
        print("Cloning header/footer styles from template...")
        
        if not template_doc.sections or not target_doc.sections:
            return
        
        template_section = template_doc.sections[0]
        target_section = target_doc.sections[0]
        
        # Clone header styles
        if template_section.header and target_section.header:
            for i, template_para in enumerate(template_section.header.paragraphs):
                if i < len(target_section.header.paragraphs):
                    target_para = target_section.header.paragraphs[i]
                    # Apply the same style
                    if template_para.style:
                        target_para.style = template_para.style
                    # Apply individual run styling
                    for j, template_run in enumerate(template_para.runs):
                        if j < len(target_para.runs):
                            target_run = target_para.runs[j]
                            if template_run.font.name:
                                target_run.font.name = template_run.font.name
                            if template_run.font.size:
                                target_run.font.size = template_run.font.size
                            if template_run.font.bold is not None:
                                target_run.font.bold = template_run.font.bold
                            if template_run.font.italic is not None:
                                target_run.font.italic = template_run.font.italic
                            if template_run.font.color.rgb:
                                target_run.font.color.rgb = template_run.font.color.rgb
        
        # Clone footer styles
        if template_section.footer and target_section.footer:
            for i, template_para in enumerate(template_section.footer.paragraphs):
                if i < len(target_section.footer.paragraphs):
                    target_para = target_section.footer.paragraphs[i]
                    # Apply the same style
                    if template_para.style:
                        target_para.style = template_para.style
                    # Apply individual run styling
                    for j, template_run in enumerate(template_para.runs):
                        if j < len(target_para.runs):
                            target_run = target_para.runs[j]
                            if template_run.font.name:
                                target_run.font.name = template_run.font.name
                            if template_run.font.size:
                                target_run.font.size = template_run.font.size
                            if template_run.font.bold is not None:
                                target_run.font.bold = template_run.font.bold
                            if template_run.font.italic is not None:
                                target_run.font.italic = template_run.font.italic
                            if template_run.font.color.rgb:
                                target_run.font.color.rgb = template_run.font.color.rgb
        
        print("Header/footer styles cloned successfully")
        
    except Exception as e:
        print(f"Warning: Could not clone header/footer styles: {e}")

# Main execution
print("Starting document generation process...")

# Check template styles first
check_template_styles(TEMPLATE_PATH)

# Load template
doc = Document(TEMPLATE_PATH)

# Clone header/footer styles from template
clone_header_footer_styles(doc, doc)

# Clear existing content
clear_document(doc)

# Parse JSON content
json_data = parse_spec_json(CONTENT_PATH)

# Apply document settings from JSON
apply_document_settings_from_json(doc, json_data)

# Apply margins from JSON
apply_margins_from_json(doc, json_data)

# Apply style definitions from JSON
apply_style_definitions_from_json(doc, json_data)

# Generate content from JSON
generate_content_from_v3_json(doc, json_data)

# Fix header and footer fonts
# The hardcoded Arial font fix is removed. Header/footer styles are now cloned.

# Save document
# Suppress the final print statement to avoid UnicodeEncodeError in batch runs
doc.save(OUTPUT_PATH)
# print(f"Document saved as '{OUTPUT_PATH}' with {FONT_SIZE}pt {FONT_NAME} font")
print(f"Content source: {CONTENT_PATH}")
print("Note: The template's multilevel list style should be applied automatically")
print("if the paragraphs use the correct style names from the template.")






