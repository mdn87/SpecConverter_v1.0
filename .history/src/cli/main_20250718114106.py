"""
Main CLI entry point for SpecConverter v1.0

Provides the command-line interface for the application.
"""

import argparse
import sys
from pathlib import Path
from typing import Optional

# Add the src directory to the path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

from utils.logging_utils import setup_logging, get_logger
from utils.file_utils import ensure_directory
from core.extractor import SpecContentExtractorV3 as SpecExtractor
from core.template_analyzer import TemplateListDetector as TemplateAnalyzer


def create_parser() -> argparse.ArgumentParser:
    """Create the command-line argument parser"""
    parser = argparse.ArgumentParser(
        prog="specconverter",
        description="SpecConverter v1.0 - Specification document conversion and processing toolkit",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Extract content from a single document
  specconverter extract document.docx --template template.docx --output output/

  # Generate a document from JSON
  specconverter generate document.json --template template.docx --output result.docx

  # Analyze a template
  specconverter template analyze template.docx

  # Process a batch job
  specconverter batch process job.yaml
        """
    )
    
    # Global options
    parser.add_argument(
        "--log-level",
        choices=["DEBUG", "INFO", "WARNING", "ERROR"],
        default="INFO",
        help="Set the logging level (default: INFO)"
    )
    
    parser.add_argument(
        "--config",
        type=str,
        help="Path to configuration file"
    )
    
    # Subcommands
    subparsers = parser.add_subparsers(dest="command", help="Available commands")
    
    # Extract command
    extract_parser = subparsers.add_parser("extract", help="Extract content from a document")
    extract_parser.add_argument("document", help="Path to the document to extract")
    extract_parser.add_argument("--template", help="Path to template file")
    extract_parser.add_argument("--output", default="output", help="Output directory")
    extract_parser.add_argument("--save-modular", action="store_true", help="Save modular JSON files")
    
    # PDF Extract command
    pdf_extract_parser = subparsers.add_parser("pdf-extract", help="Extract content using PDF conversion and OCR")
    pdf_extract_parser.add_argument("document", help="Path to the document to extract")
    pdf_extract_parser.add_argument("--template", help="Path to template file")
    pdf_extract_parser.add_argument("--output", default="output", help="Output directory")
    pdf_extract_parser.add_argument("--save-modular", action="store_true", help="Save modular JSON files")
    
    # Hybrid Analysis command
    hybrid_parser = subparsers.add_parser("hybrid", help="Hybrid analysis with PDF validation and cross-reference")
    hybrid_parser.add_argument("document", help="Path to the document to analyze")
    hybrid_parser.add_argument("--template", help="Path to template file")
    hybrid_parser.add_argument("--output", default="output", help="Output directory")
    hybrid_parser.add_argument("--save-modular", action="store_true", help="Save modular JSON files")
    hybrid_parser.add_argument("--validation-report", action="store_true", help="Generate detailed validation report")
    hybrid_parser.add_argument("--generate-word", action="store_true", help="Generate Word document with analysis results")
    
    # Generate command
    generate_parser = subparsers.add_parser("generate", help="Generate document from JSON")
    generate_parser.add_argument("json_file", help="Path to JSON file")
    generate_parser.add_argument("--template", required=True, help="Path to template file")
    generate_parser.add_argument("--output", required=True, help="Output document path")
    
    # Template command
    template_parser = subparsers.add_parser("template", help="Template management")
    template_subparsers = template_parser.add_subparsers(dest="template_command")
    
    analyze_parser = template_subparsers.add_parser("analyze", help="Analyze template structure")
    analyze_parser.add_argument("template", help="Path to template file")
    analyze_parser.add_argument("--output", help="Output file for analysis")
    
    # Batch command
    batch_parser = subparsers.add_parser("batch", help="Batch processing")
    batch_subparsers = batch_parser.add_subparsers(dest="batch_command")
    
    process_parser = batch_subparsers.add_parser("process", help="Process a batch job")
    process_parser.add_argument("job_file", help="Path to batch job configuration file")
    
    validate_parser = batch_subparsers.add_parser("validate", help="Validate a batch job")
    validate_parser.add_argument("job_file", help="Path to batch job configuration file")
    
    return parser


def extract_command(args: argparse.Namespace) -> int:
    """Handle the extract command"""
    logger = get_logger("extract")
    
    try:
        # Ensure output directory exists
        ensure_directory(args.output)
        
        # Initialize extractor
        extractor = SpecExtractor(template_path=args.template)
        
        # Extract content
        logger.info(f"Extracting content from: {args.document}")
        result = extractor.extract_content(args.document)
        
        # Save results
        base_name = Path(args.document).stem
        output_path = Path(args.output) / f"{base_name}_v3.json"
        
        # Save main JSON file
        extractor.save_to_json(result, str(output_path))
        logger.info(f"Saved main JSON to: {output_path}")
        
        # Save modular files if requested
        if args.save_modular:
            extractor.save_modular_json_files(result, base_name, args.output)
            logger.info("Saved modular JSON files")
        
        return 0
        
    except Exception as e:
        logger.error(f"Extraction failed: {e}")
        return 1


def pdf_extract_command(args: argparse.Namespace) -> int:
    """Handle the PDF extract command"""
    logger = get_logger("pdf-extract")
    
    try:
        # Import PDF extractor
        from core.pdf_extractor import extract_via_pdf
        
        # Ensure output directory exists
        ensure_directory(args.output)
        
        # Extract content using PDF conversion and OCR
        logger.info(f"Extracting content from: {args.document} using PDF conversion and OCR")
        result = extract_via_pdf(args.document, args.template)
        
        # Save results
        base_name = Path(args.document).stem
        output_path = Path(args.output) / f"{base_name}_pdf_extract.json"
        
        # Convert to JSON format
        import json
        json_data = {
            "file_path": result.file_path,
            "content_blocks": [
                {
                    "text": block.text,
                    "level_type": block.level_type,
                    "number": block.number,
                    "content": block.content,
                    "level_number": block.level_number,
                    "bwa_level_name": block.bwa_level_name,
                    "numbering_id": block.numbering_id,
                    "numbering_level": block.numbering_level,
                    "style_name": block.style_name
                }
                for block in result.content_blocks
            ],
            "header_footer": {
                "header": result.header_footer.header,
                "footer": result.header_footer.footer,
                "margins": result.header_footer.margins,
                "document_settings": result.header_footer.document_settings
            },
            "template_analysis": {
                "template_path": result.template_analysis.template_path if result.template_analysis else "",
                "analysis_timestamp": result.template_analysis.analysis_timestamp if result.template_analysis else "",
                "numbering_definitions": result.template_analysis.numbering_definitions if result.template_analysis else {},
                "bwa_list_levels": result.template_analysis.bwa_list_levels if result.template_analysis else {},
                "level_mappings": result.template_analysis.level_mappings if result.template_analysis else {},
                "summary": result.template_analysis.summary if result.template_analysis else {}
            },
            "validation_results": {
                "errors": [{"line_number": e.line_number, "error_type": e.error_type, "message": e.message, "context": e.context} for e in result.validation_results.errors],
                "corrections": result.validation_results.corrections,
                "validation_summary": result.validation_results.validation_summary
            }
        }
        
        # Save main JSON file
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Saved PDF extraction results to: {output_path}")
        
        # Save modular files if requested
        if args.save_modular:
            # Create modular files directory
            modular_dir = Path(args.output) / f"{base_name}_pdf_modular"
            modular_dir.mkdir(exist_ok=True)
            
            # Save content blocks separately
            for i, block in enumerate(result.content_blocks):
                block_name = block.number.replace(' ', '_') if block.number else f"block_{i+1:03d}"
                block_file = modular_dir / f"block_{i+1:03d}_{block_name}.json"
                with open(block_file, 'w', encoding='utf-8') as f:
                    json.dump({
                        "text": block.text,
                        "level_type": block.level_type,
                        "number": block.number,
                        "content": block.content,
                        "level_number": block.level_number,
                        "bwa_level_name": block.bwa_level_name,
                        "numbering_id": block.numbering_id,
                        "numbering_level": block.numbering_level,
                        "style_name": block.style_name
                    }, f, indent=2, ensure_ascii=False)
            
            logger.info(f"Saved modular files to: {modular_dir}")
        
        return 0
        
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return 1


def hybrid_command(args: argparse.Namespace) -> int:
    """Handle the hybrid analysis command"""
    logger = get_logger("hybrid")
    
    try:
        # Import hybrid analyzer
        from core.hybrid_analyzer import analyze_with_hybrid_validation
        
        # Ensure output directory exists
        ensure_directory(args.output)
        
        # Perform hybrid analysis
        logger.info(f"Starting hybrid analysis of: {args.document}")
        result = analyze_with_hybrid_validation(args.document, args.template)
        
        # Save results
        base_name = Path(args.document).stem
        output_path = Path(args.output) / f"{base_name}_hybrid_analysis.json"
        
        # Convert to JSON format
        import json
        json_data = {
            "file_path": result.file_path,
            "content_blocks": [
                {
                    "text": block.text,
                    "level_type": block.level_type,
                    "number": block.number,
                    "content": block.content,
                    "level_number": block.level_number,
                    "bwa_level_name": block.bwa_level_name,
                    "numbering_id": block.numbering_id,
                    "numbering_level": block.numbering_level,
                    "style_name": block.style_name
                }
                for block in result.content_blocks
            ],
            "header_footer": {
                "header": result.header_footer.header,
                "footer": result.header_footer.footer,
                "margins": result.header_footer.margins,
                "document_settings": result.header_footer.document_settings
            },
            "template_analysis": {
                "template_path": result.template_analysis.template_path if result.template_analysis else "",
                "analysis_timestamp": result.template_analysis.analysis_timestamp if result.template_analysis else "",
                "numbering_definitions": result.template_analysis.numbering_definitions if result.template_analysis else {},
                "bwa_list_levels": result.template_analysis.bwa_list_levels if result.template_analysis else {},
                "level_mappings": result.template_analysis.level_mappings if result.template_analysis else {},
                "summary": result.template_analysis.summary if result.template_analysis else {}
            },
            "validation_results": {
                "errors": [{"line_number": e.line_number, "error_type": e.error_type, "message": e.message, "context": e.context} for e in result.validation_results.errors],
                "corrections": result.validation_results.corrections,
                "validation_summary": result.validation_results.validation_summary
            }
        }
        
        # Save main JSON file
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Saved hybrid analysis results to: {output_path}")
        
        # Generate validation report if requested
        if args.validation_report:
            from core.hybrid_analyzer import HybridAnalyzer
            analyzer = HybridAnalyzer(args.template)
            analyzer.analyze_document(args.document, args.template)
            report = analyzer.get_validation_report()
            
            report_path = Path(args.output) / f"{base_name}_validation_report.json"
            with open(report_path, 'w', encoding='utf-8') as f:
                json.dump(report, f, indent=2, ensure_ascii=False)
            
            logger.info(f"Saved validation report to: {report_path}")
            
            # Print summary
            print(f"\n=== Hybrid Analysis Validation Report ===")
            print(f"Total blocks processed: {report['summary']['total_blocks_processed']}")
            print(f"Blocks with numbering corrections: {report['summary']['blocks_with_numbering_corrections']}")
            print(f"Blocks not found in PDF: {report['summary']['blocks_not_found_in_pdf']}")
            print(f"PDF content length: {report['pdf_content_length']} characters")
            print(f"Template patterns found: {report['template_patterns']}")
        
        # Save modular files if requested
        if args.save_modular:
            # Create modular files directory
            modular_dir = Path(args.output) / f"{base_name}_hybrid_modular"
            modular_dir.mkdir(exist_ok=True)
            
            # Save content blocks separately
            for i, block in enumerate(result.content_blocks):
                block_name = block.number.replace(' ', '_') if block.number else f"block_{i+1:03d}"
                block_file = modular_dir / f"block_{i+1:03d}_{block_name}.json"
                with open(block_file, 'w', encoding='utf-8') as f:
                    json.dump({
                        "text": block.text,
                        "level_type": block.level_type,
                        "number": block.number,
                        "content": block.content,
                        "level_number": block.level_number,
                        "bwa_level_name": block.bwa_level_name,
                        "numbering_id": block.numbering_id,
                        "numbering_level": block.numbering_level,
                        "style_name": block.style_name
                    }, f, indent=2, ensure_ascii=False)
            
            logger.info(f"Saved modular files to: {modular_dir}")
        
        # Generate Word document if requested
        if args.generate_word:
            logger.info("Generating Word document with analysis results...")
            word_path = Path(args.output) / f"{base_name}_hybrid_analysis_report.docx"
            
            # Create Word document with analysis results
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title
            title = doc.add_heading('Hybrid Analysis Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Document information
            doc.add_heading('Document Information', level=1)
            doc.add_paragraph(f"Source Document: {args.document}")
            doc.add_paragraph(f"Template: {args.template or 'None'}")
            doc.add_paragraph(f"Analysis Date: {result.template_analysis.analysis_timestamp if result.template_analysis else 'N/A'}")
            
            # Summary statistics
            doc.add_heading('Analysis Summary', level=1)
            summary_table = doc.add_table(rows=1, cols=2)
            summary_table.style = 'Table Grid'
            hdr_cells = summary_table.rows[0].cells
            hdr_cells[0].text = 'Metric'
            hdr_cells[1].text = 'Value'
            
            # Add summary data
            summary_data = [
                ('Total Content Blocks', str(len(result.content_blocks))),
                ('PDF Content Length', f"{report['pdf_content_length'] if args.validation_report else 'N/A'} characters"),
                ('Template Patterns Found', f"{report['template_patterns'] if args.validation_report else 'N/A'}"),
                ('Numbering Corrections', f"{report['summary']['blocks_with_numbering_corrections'] if args.validation_report else 'N/A'}"),
                ('Blocks Not Found in PDF', f"{report['summary']['blocks_not_found_in_pdf'] if args.validation_report else 'N/A'}")
            ]
            
            for metric, value in summary_data:
                row_cells = summary_table.add_row().cells
                row_cells[0].text = metric
                row_cells[1].text = value
            
            # Content blocks analysis
            doc.add_heading('Content Blocks Analysis', level=1)
            
            # Create table for content blocks
            blocks_table = doc.add_table(rows=1, cols=5)
            blocks_table.style = 'Table Grid'
            hdr_cells = blocks_table.rows[0].cells
            hdr_cells[0].text = 'Block #'
            hdr_cells[1].text = 'Number'
            hdr_cells[2].text = 'Level'
            hdr_cells[3].text = 'Type'
            hdr_cells[4].text = 'Text Preview'
            
            # Add content blocks data
            for i, block in enumerate(result.content_blocks[:20]):  # Limit to first 20 blocks
                row_cells = blocks_table.add_row().cells
                row_cells[0].text = str(i + 1)
                row_cells[1].text = block.number or 'N/A'
                row_cells[2].text = str(block.level_number) if block.level_number is not None else 'N/A'
                row_cells[3].text = block.level_type or 'N/A'
                row_cells[4].text = (block.text[:50] + '...') if len(block.text) > 50 else block.text
            
            if len(result.content_blocks) > 20:
                doc.add_paragraph(f"... and {len(result.content_blocks) - 20} more blocks")
            
            # Template analysis
            if result.template_analysis and result.template_analysis.bwa_list_levels:
                doc.add_heading('Template Analysis', level=1)
                doc.add_paragraph(f"Template Path: {result.template_analysis.template_path}")
                doc.add_paragraph(f"BWA List Levels Found: {len(result.template_analysis.bwa_list_levels)}")
                
                # Template levels table
                levels_table = doc.add_table(rows=1, cols=3)
                levels_table.style = 'Table Grid'
                hdr_cells = levels_table.rows[0].cells
                hdr_cells[0].text = 'Level'
                hdr_cells[1].text = 'BWA Style'
                hdr_cells[2].text = 'Format'
                
                for level, level_info in result.template_analysis.bwa_list_levels.items():
                    row_cells = levels_table.add_row().cells
                    row_cells[0].text = str(level)
                    row_cells[1].text = level_info.get('style_name', 'N/A')
                    row_cells[2].text = level_info.get('format', 'N/A')
            
            # Validation results
            if result.validation_results and result.validation_results.errors:
                doc.add_heading('Validation Issues', level=1)
                for error in result.validation_results.errors:
                    doc.add_paragraph(f"Error: {error.message}", style='List Bullet')
            
            # Save the document
            doc.save(str(word_path))
            logger.info(f"Generated Word report: {word_path}")
        
        return 0
        
    except Exception as e:
        logger.error(f"Hybrid analysis failed: {e}")
        return 1


def generate_command(args: argparse.Namespace) -> int:
    """Handle the generate command"""
    logger = get_logger("generate")
    
    try:
        # Import generator functions
        from core.generator import generate_content_from_v3_json, parse_spec_json, apply_document_settings_from_json, apply_margins_from_json, apply_style_definitions_from_json, clear_document, clone_header_footer_styles
        
        # Load JSON data
        logger.info(f"Loading JSON data from: {args.json_file}")
        json_data = parse_spec_json(args.json_file)
        
        if not json_data:
            logger.error("Failed to load JSON data")
            return 1
        
        # Load template
        from docx import Document
        logger.info(f"Loading template: {args.template}")
        doc = Document(args.template)
        
        # Clone header/footer styles from template
        clone_header_footer_styles(doc, doc)
        
        # Clear existing content
        clear_document(doc)
        
        # Apply document settings from JSON
        apply_document_settings_from_json(doc, json_data)
        
        # Apply margins from JSON
        apply_margins_from_json(doc, json_data)
        
        # Apply style definitions from JSON
        apply_style_definitions_from_json(doc, json_data)
        
        # Generate content
        logger.info(f"Generating document from: {args.json_file}")
        generate_content_from_v3_json(doc, json_data)
        
        # Save document
        doc.save(args.output)
        
        logger.info(f"Generated document: {args.output}")
        return 0
        
    except Exception as e:
        logger.error(f"Generation failed: {e}")
        return 1


def template_analyze_command(args: argparse.Namespace) -> int:
    """Handle the template analyze command"""
    logger = get_logger("template")
    
    try:
        # Initialize analyzer
        analyzer = TemplateAnalyzer()
        
        # Analyze template
        logger.info(f"Analyzing template: {args.template}")
        analysis = analyzer.analyze_template(args.template)
        
        # Save or display results
        if args.output:
            analyzer.save_analysis_to_json(analysis, args.output)
            logger.info(f"Analysis saved to: {args.output}")
        else:
            print(f"Template Analysis for: {analysis.template_path}")
            print(f"BWA List Levels: {len(analysis.bwa_list_levels)}")
            print(f"Numbering Definitions: {len(analysis.numbering_definitions)}")
        
        return 0
        
    except Exception as e:
        logger.error(f"Template analysis failed: {e}")
        return 1


def batch_process_command(args: argparse.Namespace) -> int:
    """Handle the batch process command"""
    logger = get_logger("batch")
    
    try:
        # Import here to avoid circular imports
        from batch.processor import BatchProcessor
        
        # Initialize processor
        processor = BatchProcessor()
        
        # Process batch job
        logger.info(f"Processing batch job: {args.job_file}")
        results = processor.process_job(args.job_file)
        
        # Display results
        print(f"Batch processing complete:")
        print(f"  Total processed: {results.total_processed}")
        print(f"  Successful: {results.total_successful}")
        print(f"  Failed: {results.total_failed}")
        
        return 0 if results.total_failed == 0 else 1
        
    except Exception as e:
        logger.error(f"Batch processing failed: {e}")
        return 1


def batch_validate_command(args: argparse.Namespace) -> int:
    """Handle the batch validate command"""
    logger = get_logger("batch")
    
    try:
        # Import here to avoid circular imports
        from batch.processor import BatchProcessor
        
        # Initialize processor
        processor = BatchProcessor()
        
        # Validate batch job
        logger.info(f"Validating batch job: {args.job_file}")
        is_valid = processor.validate_job_config(args.job_file)
        
        if is_valid:
            print("✓ Batch job configuration is valid")
            return 0
        else:
            print("✗ Batch job configuration has errors")
            return 1
        
    except Exception as e:
        logger.error(f"Batch validation failed: {e}")
        return 1


def main() -> int:
    """Main entry point"""
    parser = create_parser()
    args = parser.parse_args()
    
    # Set up logging
    setup_logging(level=args.log_level)
    logger = get_logger("main")
    
    # Handle commands
    if args.command == "extract":
        return extract_command(args)
    elif args.command == "pdf-extract":
        return pdf_extract_command(args)
    elif args.command == "hybrid":
        return hybrid_command(args)
    elif args.command == "generate":
        return generate_command(args)
    elif args.command == "template" and args.template_command == "analyze":
        return template_analyze_command(args)
    elif args.command == "batch" and args.batch_command == "process":
        return batch_process_command(args)
    elif args.command == "batch" and args.batch_command == "validate":
        return batch_validate_command(args)
    else:
        parser.print_help()
        return 1


if __name__ == "__main__":
    sys.exit(main()) 