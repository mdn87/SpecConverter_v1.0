#!/usr/bin/env python3
"""
Header and Footer Extractor Module

This module provides functionality to extract header, footer, margin, and comment information
from Word documents (.docx). It combines functionality from the rip scripts into a reusable module.

Features:
- Extract header content (paragraphs, tables, text boxes)
- Extract footer content (paragraphs, tables, text boxes)
- Extract margin settings
- Extract comments with metadata
- Save data to JSON and TXT formats

Usage:
    from header_footer_extractor import HeaderFooterExtractor
    
    extractor = HeaderFooterExtractor()
    data = extractor.extract_all(docx_path)
"""

from docx import Document
import json
import os
from typing import Dict, List, Any, Optional
from datetime import datetime

class HeaderFooterExtractor:
    """Extracts header, footer, margin, and comment information from Word documents"""
    
    def __init__(self):
        pass
    
    def extract_all(self, docx_path: str) -> Dict[str, Any]:
        """
        Extract all header, footer, margin, and comment information from a document
        
        Args:
            docx_path: Path to the Word document
            
        Returns:
            Dictionary containing header, footer, margin, and comment data
        """
        try:
            doc = Document(docx_path)
            
            # Extract all components
            header_footer_data = self.extract_header_footer_margins(docx_path)
            comments = self.extract_comments(docx_path)
            
            return {
                "header": header_footer_data["header"],
                "footer": header_footer_data["footer"],
                "margins": header_footer_data["margins"],
                "comments": comments
            }
            
        except Exception as e:
            print(f"Error extracting header/footer data: {e}")
            return {
                "header": {"paragraphs": [], "tables": [], "text_boxes": []},
                "footer": {"paragraphs": [], "tables": [], "text_boxes": []},
                "margins": {},
                "comments": []
            }
    
    def extract_header_footer_margins(self, docx_path: str) -> Dict[str, Any]:
        """
        Extract header, footer, margin, and document-level settings from a Word document
        
        Args:
            docx_path: Path to the Word document
            
        Returns:
            Dictionary containing header, footer, margin, and document settings
        """
        try:
            doc = Document(docx_path)
            sec = doc.sections[0]
            
            # Extract margin settings
            margins = self._extract_margins(sec)
            
            # Extract document-level settings
            document_settings = self._extract_document_settings(sec, doc)
            
            # Extract header content
            header_content = self._extract_header_content(sec)
            
            # Extract footer content
            footer_content = self._extract_footer_content(sec)
            
            return {
                "header": header_content,
                "footer": footer_content,
                "margins": margins,
                "document_settings": document_settings
            }
            
        except Exception as e:
            print(f"Error extracting header/footer/margins: {e}")
            return {
                "header": {"paragraphs": [], "tables": [], "text_boxes": []},
                "footer": {"paragraphs": [], "tables": [], "text_boxes": []},
                "margins": {},
                "document_settings": {}
            }
    
    def _extract_margins(self, section: Any) -> Dict[str, float]:
        """Extract margin settings from a document section"""
        margins = {}
        try:
            if section.top_margin:
                margins["top_margin"] = section.top_margin.inches
            if section.bottom_margin:
                margins["bottom_margin"] = section.bottom_margin.inches
            if section.left_margin:
                margins["left_margin"] = section.left_margin.inches
            if section.right_margin:
                margins["right_margin"] = section.right_margin.inches
            if section.header_distance:
                margins["header_distance"] = section.header_distance.inches
            if section.footer_distance:
                margins["footer_distance"] = section.footer_distance.inches
        except Exception as e:
            print(f"Warning: Could not extract margin settings: {e}")
        
        return margins
    
    def _extract_document_settings(self, section: Any, doc: Any) -> Dict[str, Any]:
        """Extract comprehensive document-level settings"""
        settings = {}
        
        try:
            # Page size and orientation
            if section.page_width:
                settings["page_width"] = section.page_width.inches
            if section.page_height:
                settings["page_height"] = section.page_height.inches
            
            # Page orientation
            if section.page_width and section.page_height:
                if section.page_width.inches > section.page_height.inches:
                    settings["page_orientation"] = "landscape"
                else:
                    settings["page_orientation"] = "portrait"
            
            # Page margins (already in margins, but keeping for completeness)
            if section.top_margin:
                settings["top_margin"] = section.top_margin.inches
            if section.bottom_margin:
                settings["bottom_margin"] = section.bottom_margin.inches
            if section.left_margin:
                settings["left_margin"] = section.left_margin.inches
            if section.right_margin:
                settings["right_margin"] = section.right_margin.inches
            
            # Header and footer distances
            if section.header_distance:
                settings["header_distance"] = section.header_distance.inches
            if section.footer_distance:
                settings["footer_distance"] = section.footer_distance.inches
            
            # Gutter settings
            if section.gutter:
                settings["gutter"] = section.gutter.inches
            
            # Different first page header/footer
            settings["different_first_page_header_footer"] = section.different_first_page_header_footer
            
            # Different odd and even pages
            try:
                settings["different_odd_and_even_pages"] = section.different_odd_and_even_pages
            except AttributeError:
                # Some versions of python-docx don't have this attribute
                settings["different_odd_and_even_pages"] = None
            
            # Page numbering
            if hasattr(section, 'page_numbering') and section.page_numbering:
                settings["page_numbering"] = {
                    "start": section.page_numbering.start,
                    "restart": section.page_numbering.restart,
                    "format": section.page_numbering.format
                }
            
            # Line numbering
            if hasattr(section, 'line_numbering') and section.line_numbering:
                settings["line_numbering"] = {
                    "start": section.line_numbering.start,
                    "increment": section.line_numbering.increment,
                    "restart": section.line_numbering.restart,
                    "distance": section.line_numbering.distance.inches if section.line_numbering.distance else None
                }
            
            # Document properties
            if doc.core_properties:
                settings["document_properties"] = {
                    "title": doc.core_properties.title,
                    "subject": doc.core_properties.subject,
                    "author": doc.core_properties.author,
                    "keywords": doc.core_properties.keywords,
                    "category": doc.core_properties.category,
                    "comments": doc.core_properties.comments,
                    "created": doc.core_properties.created.isoformat() if doc.core_properties.created else None,
                    "modified": doc.core_properties.modified.isoformat() if doc.core_properties.modified else None,
                    "last_modified_by": doc.core_properties.last_modified_by,
                    "revision": doc.core_properties.revision
                }
            
            # Extract default paragraph and run formatting from styles
            settings["default_formatting"] = self._extract_default_formatting(doc)
            
            # Extract document-wide settings from settings.xml
            settings["document_wide_settings"] = self._extract_document_wide_settings(doc)
            
        except Exception as e:
            print(f"Warning: Could not extract document settings: {e}")
        
        return settings
    
    def _extract_default_formatting(self, doc: Any) -> Dict[str, Any]:
        """Extract default formatting settings from document styles"""
        default_formatting = {}
        
        try:
            # Get the Normal style
            normal_style = doc.styles.get('Normal')
            if normal_style:
                # Default paragraph format
                if normal_style.paragraph_format:
                    pf = normal_style.paragraph_format
                    default_formatting["default_paragraph_format"] = {
                        "alignment": str(pf.alignment) if pf.alignment else None,
                        "left_indent": pf.left_indent.inches if pf.left_indent else None,
                        "right_indent": pf.right_indent.inches if pf.right_indent else None,
                        "first_line_indent": pf.first_line_indent.inches if pf.first_line_indent else None,
                        "space_before": pf.space_before.pt if pf.space_before else None,
                        "space_after": pf.space_after.pt if pf.space_after else None,
                        "line_spacing": pf.line_spacing,
                        "keep_with_next": pf.keep_with_next,
                        "keep_lines_together": pf.keep_lines_together,
                        "page_break_before": pf.page_break_before,
                        "widow_control": pf.widow_control
                    }
                
                # Default run format (font)
                if normal_style.font:
                    font = normal_style.font
                    default_formatting["default_run_format"] = {
                        "name": font.name,
                        "size": font.size.pt if font.size else None,
                        "bold": font.bold,
                        "italic": font.italic,
                        "underline": str(font.underline) if font.underline else None,
                        "color": str(font.color.rgb) if font.color.rgb else None,
                        "strike": font.strike,
                        "small_caps": font.small_caps,
                        "all_caps": font.all_caps
                    }
                    
        except Exception as e:
            print(f"Warning: Could not extract default formatting: {e}")
        
        return default_formatting
    
    def _extract_document_wide_settings(self, doc: Any) -> Dict[str, Any]:
        """Extract document-wide settings from settings.xml"""
        doc_settings = {}
        
        try:
            # Access the settings.xml through the document's _element
            settings_element = doc._element.find('.//w:settings', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if settings_element is not None:
                # Extract various document-wide settings
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                
                # Default tab stop
                try:
                    default_tab_stop = settings_element.find('.//w:defaultTabStop', ns)
                    if default_tab_stop is not None:
                        doc_settings["default_tab_stop"] = default_tab_stop.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                except:
                    pass
                
                # Character spacing control
                try:
                    char_spacing_control = settings_element.find('.//w:characterSpacingControl', ns)
                    if char_spacing_control is not None:
                        doc_settings["character_spacing_control"] = char_spacing_control.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                except:
                    pass
                
                # Compatibility settings
                try:
                    compatibility = settings_element.find('.//w:compat', ns)
                    if compatibility is not None:
                        doc_settings["compatibility_settings"] = {}
                        for child in compatibility:
                            try:
                                tag = child.tag.split('}')[-1]  # Remove namespace
                                doc_settings["compatibility_settings"][tag] = child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                            except:
                                continue
                except:
                    pass
                
                # Document protection
                try:
                    document_protection = settings_element.find('.//w:documentProtection', ns)
                    if document_protection is not None:
                        doc_settings["document_protection"] = {
                            "enforcement": document_protection.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}enforcement'),
                            "edit": document_protection.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}edit'),
                            "formatting": document_protection.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}formatting')
                        }
                except:
                    pass
                
                # Zoom settings
                try:
                    zoom = settings_element.find('.//w:zoom', ns)
                    if zoom is not None:
                        doc_settings["zoom"] = {
                            "percent": zoom.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}percent'),
                            "val": zoom.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                        }
                except:
                    pass
                
                # View settings
                try:
                    view = settings_element.find('.//w:view', ns)
                    if view is not None:
                        doc_settings["view"] = {
                            "val": view.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'),
                            "zoom": view.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}zoom')
                        }
                except:
                    pass
                
                # Grammar and spelling settings
                try:
                    proof_state = settings_element.find('.//w:proofState', ns)
                    if proof_state is not None:
                        doc_settings["proof_state"] = {
                            "grammar": proof_state.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}grammar'),
                            "spelling": proof_state.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spelling')
                        }
                except:
                    pass
                
                # Track changes settings
                try:
                    track_changes = settings_element.find('.//w:trackRevisions', ns)
                    if track_changes is not None:
                        doc_settings["track_changes"] = track_changes.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                except:
                    pass
                
                # Print settings
                try:
                    print_settings = settings_element.find('.//w:printSettings', ns)
                    if print_settings is not None:
                        doc_settings["print_settings"] = {}
                        for child in print_settings:
                            try:
                                tag = child.tag.split('}')[-1]
                                doc_settings["print_settings"][tag] = child.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                            except:
                                continue
                except:
                    pass
            
        except Exception as e:
            # Silently handle any extraction errors
            pass
        
        return doc_settings
    
    def _extract_header_content(self, section: Any) -> Dict[str, List]:
        """Extract header content from a document section"""
        return self._extract_content_from_section(section.header._element, section.header._element.nsmap)
    
    def _extract_footer_content(self, section: Any) -> Dict[str, List]:
        """Extract footer content from a document section"""
        return self._extract_content_from_section(section.footer._element, section.footer._element.nsmap)
    
    def _extract_content_from_section(self, section_element: Any, nsmap: Dict[str, str]) -> Dict[str, List]:
        """
        Extract content from header or footer section
        
        Args:
            section_element: XML element representing the section
            nsmap: Namespace mapping for XML parsing
            
        Returns:
            Dictionary containing paragraphs, tables, and text boxes
        """
        content = {
            "paragraphs": [],
            "tables": [],
            "text_boxes": []
        }
        
        # Extract paragraphs
        for p in section_element.findall('.//w:p', namespaces=nsmap):
            text = self._extract_text_from_element(p, nsmap)
            if text:
                content["paragraphs"].append(text)
        
        # Extract tables
        for tbl in section_element.findall('.//w:tbl', namespaces=nsmap):
            table_data = []
            for row in tbl.findall('.//w:tr', namespaces=nsmap):
                row_data = []
                for cell in row.findall('.//w:tc', namespaces=nsmap):
                    cell_text = self._extract_text_from_element(cell, nsmap)
                    row_data.append(cell_text)
                if row_data:
                    table_data.append(row_data)
            if table_data:
                content["tables"].append(table_data)
        
        # Extract text boxes
        for drawing in section_element.findall('.//w:txbxContent', namespaces=nsmap):
            textbox_data = []
            for p in drawing.findall('.//w:p', namespaces=nsmap):
                text = self._extract_text_from_element(p, nsmap)
                if text:
                    textbox_data.append(text)
            if textbox_data:
                content["text_boxes"].append(textbox_data)
        
        return content
    
    def _extract_text_from_element(self, element: Any, nsmap: Dict[str, str]) -> str:
        """Extract text content from an XML element"""
        text_parts = []
        
        # Extract text from the element itself
        if element.text:
            text_parts.append(element.text.strip())
        
        # Extract text from child elements
        for child in element:
            if child.text:
                text_parts.append(child.text.strip())
            if child.tail:
                text_parts.append(child.tail.strip())
        
        return " ".join(text_parts).strip()
    
    def extract_comments(self, docx_path: str) -> List[Dict[str, Any]]:
        """
        Extract comments from a Word document
        
        Args:
            docx_path: Path to the Word document
            
        Returns:
            List of comment dictionaries with metadata
        """
        try:
            doc = Document(docx_path)
            comments = []
            
            # Check if the document has comments
            if hasattr(doc.part, '_comments_part') and doc.part._comments_part is not None:
                for c in doc.part._comments_part.comments:
                    # Assemble the full comment text
                    full_text = "\n".join(p.text for p in c.paragraphs).strip()
                    
                    comment_data = {
                        "text": full_text,
                        "ref": None
                    }
                    
                    # Get comment metadata
                    try:
                        comment_data["author"] = str(c.author) if c.author else None
                    except:
                        comment_data["author"] = None
                    
                    try:
                        comment_data["timestamp"] = str(c.timestamp) if c.timestamp else None
                    except:
                        comment_data["timestamp"] = None
                    
                    try:
                        comment_data["comment_id"] = str(c.comment_id) if c.comment_id else None
                    except:
                        comment_data["comment_id"] = None
                    
                    try:
                        comment_data["initials"] = str(c.initials) if c.initials else None
                    except:
                        comment_data["initials"] = None
                    
                    comments.append(comment_data)
            else:
                print("No comments found in the document")
            
            return comments
            
        except Exception as e:
            print(f"Error extracting comments: {e}")
            return []
    
    def save_to_json(self, data: Dict[str, Any], output_path: str):
        """
        Save header/footer data to JSON file
        
        Args:
            data: Dictionary containing header/footer data
            output_path: Path to save the JSON file
        """
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        print(f"Header/footer data saved to JSON: {output_path}")
    
    def save_to_txt(self, data: Dict[str, Any], output_path: str):
        """
        Save header/footer data to TXT file
        
        Args:
            data: Dictionary containing header/footer data
            output_path: Path to save the TXT file
        """
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("HEADER, FOOTER, AND MARGIN DATA FROM DOCUMENT\n")
            f.write("=" * 60 + "\n\n")
            
            # Header information
            f.write("HEADER CONTENT:\n")
            f.write("-" * 20 + "\n")
            if data.get('header'):
                header = data['header']
                if header.get('paragraphs'):
                    f.write("Paragraphs:\n")
                    for i, para in enumerate(header['paragraphs'], 1):
                        f.write(f"  {i}. {para}\n")
                
                if header.get('tables'):
                    f.write("\nTables:\n")
                    for i, table in enumerate(header['tables'], 1):
                        f.write(f"  Table {i}:\n")
                        for j, row in enumerate(table, 1):
                            f.write(f"    Row {j}: {row}\n")
                
                if header.get('text_boxes'):
                    f.write("\nText Boxes:\n")
                    for i, textbox in enumerate(header['text_boxes'], 1):
                        f.write(f"  Text Box {i}:\n")
                        for j, para in enumerate(textbox, 1):
                            f.write(f"    {j}. {para}\n")
            else:
                f.write("No header content found\n")
            
            # Footer information
            f.write("\n\nFOOTER CONTENT:\n")
            f.write("-" * 20 + "\n")
            if data.get('footer'):
                footer = data['footer']
                if footer.get('paragraphs'):
                    f.write("Paragraphs:\n")
                    for i, para in enumerate(footer['paragraphs'], 1):
                        f.write(f"  {i}. {para}\n")
                
                if footer.get('tables'):
                    f.write("\nTables:\n")
                    for i, table in enumerate(footer['tables'], 1):
                        f.write(f"  Table {i}:\n")
                        for j, row in enumerate(table, 1):
                            f.write(f"    Row {j}: {row}\n")
                
                if footer.get('text_boxes'):
                    f.write("\nText Boxes:\n")
                    for i, textbox in enumerate(footer['text_boxes'], 1):
                        f.write(f"  Text Box {i}:\n")
                        for j, para in enumerate(textbox, 1):
                            f.write(f"    {j}. {para}\n")
            else:
                f.write("No footer content found\n")
            
            # Margin settings
            f.write("\n\nMARGIN SETTINGS:\n")
            f.write("-" * 20 + "\n")
            if data.get('margins'):
                for key, value in data['margins'].items():
                    f.write(f"{key.replace('_', ' ').title()}: {value} inches\n")
            else:
                f.write("No margin information found\n")
            
            # Comments
            f.write("\n\nCOMMENTS:\n")
            f.write("-" * 20 + "\n")
            if data.get('comments'):
                for i, comment in enumerate(data['comments'], 1):
                    f.write(f"Comment {i}:\n")
                    f.write(f"  Text: {comment['text']}\n")
                    if comment.get('author'):
                        f.write(f"  Author: {comment['author']}\n")
                    if comment.get('timestamp'):
                        f.write(f"  Timestamp: {comment['timestamp']}\n")
                    if comment.get('initials'):
                        f.write(f"  Initials: {comment['initials']}\n")
                    if comment.get('comment_id'):
                        f.write(f"  Comment ID: {comment['comment_id']}\n")
                    f.write("-" * 30 + "\n")
            else:
                f.write("No comments found\n")
        
        print(f"Header/footer data saved to TXT: {output_path}")

def main():
    """Main function for standalone usage"""
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python header_footer_extractor.py <docx_file>")
        print("Example: python header_footer_extractor.py 'SECTION 26 05 00.docx'")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    
    if not os.path.exists(docx_file):
        print(f"Error: File '{docx_file}' not found.")
        sys.exit(1)
    
    try:
        extractor = HeaderFooterExtractor()
        data = extractor.extract_all(docx_file)
        
        if data:
            # Generate output filenames
            base_name = os.path.splitext(os.path.basename(docx_file))[0]
            json_file = f"{base_name}_header_footer.json"
            txt_file = f"{base_name}_header_footer.txt"
            
            # Save to JSON
            extractor.save_to_json(data, json_file)
            
            # Save to TXT
            extractor.save_to_txt(data, txt_file)
            
            print(f"\nExtracted header, footer, and margin data from {docx_file}")
            print("Files created:")
            print(f"  - {json_file}")
            print(f"  - {txt_file}")
            
            # Show a preview of the extracted data
            print("\nData Preview:")
            
            # Header preview
            if data.get('header'):
                header = data['header']
                if header.get('paragraphs'):
                    print(f"  Header paragraphs: {len(header['paragraphs'])}")
                    for i, para in enumerate(header['paragraphs'][:2], 1):
                        print(f"    {i}. {para[:50]}...")
            
            # Footer preview
            if data.get('footer'):
                footer = data['footer']
                if footer.get('paragraphs'):
                    print(f"  Footer paragraphs: {len(footer['paragraphs'])}")
                    for i, para in enumerate(footer['paragraphs'][:2], 1):
                        print(f"    {i}. {para[:50]}...")
            
            # Comments preview
            if data.get('comments'):
                print(f"  Comments: {len(data['comments'])}")
                for i, comment in enumerate(data['comments'][:2], 1):
                    print(f"    {i}. {comment['text'][:50]}...")
            
            # Margins preview
            if data.get('margins'):
                print(f"  Margins: {len(data['margins'])} settings extracted")
        else:
            print("No data extracted")
            
    except Exception as e:
        print(f"Error processing document: {e}")

if __name__ == "__main__":
    main() 